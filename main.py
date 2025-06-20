import gradio as gr
import os
import numpy as np
from embedding import generate_embeddings, save_to_faiss
import docx
import faiss
import shutil  # 添加导入
from dotenv import load_dotenv
import fitz  # PyMuPDF

load_dotenv()  # 加载.env文件中的ZHIPU_API_KEY

# 文件存储配置
BASE_DIR = "E:/pycode/法律助手"  # 根据实际路径调整
UPLOAD_DIR = os.path.join(BASE_DIR, "uploaded_files")
EMBEDDING_DIR = os.path.join(BASE_DIR, "embeddings")

# 确保目录存在
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EMBEDDING_DIR, exist_ok=True)

# 定义向量维度
dimension = 768  

# 简化初始化逻辑
def init_system():
    global index, file_registry
    # 使用支持ID管理的索引类型
    index = faiss.IndexIDMap(faiss.IndexFlatL2(dimension))
    file_registry = {}
    
    for f in os.listdir(UPLOAD_DIR):
        if f.endswith(('.txt', '.docx', '.pdf')):
            name = os.path.splitext(f)[0]
            file_path = os.path.join(UPLOAD_DIR, f)
            embed_path = os.path.join(EMBEDDING_DIR, f"{name}_embeddings.npy")
            
            if os.path.exists(embed_path):
                embeddings = np.load(embed_path)
                # 生成唯一ID（使用文件名哈希）
                unique_id = hash(name) & 0x7FFFFFFF  # 生成正整数
                index.add_with_ids(embeddings, np.array([unique_id], dtype=np.int64))
                file_registry[name] = {
                    "path": file_path,
                    "embedding": embed_path,
                    "index_ids": [unique_id]
                }

# 调整函数定义顺序
def read_file_content(file_path):
    if file_path.endswith('.docx'):
        doc = docx.Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs] + 
                        [cell.text for table in doc.tables for row in table.rows for cell in row.cells])
    elif file_path.endswith('.pdf'):
        return ''.join([page.get_text() for page in fitz.open(file_path)])
    else:
        for enc in ['utf-8', 'gbk', 'gb18030', 'big5', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

def process_upload(file_obj, custom_name):  # 正确定义process_upload
    file_ext = os.path.splitext(file_obj.name)[1]
    file_path = os.path.join(UPLOAD_DIR, f"{custom_name}{file_ext}")
    shutil.move(file_obj.name, file_path)
    
    text = read_file_content(file_path)
    
    if not text.strip() or any(ord(c) < 32 and c not in '\t\n\r' for c in text):
        os.remove(file_path)
        raise gr.Error("文件内容无效或包含二进制数据")
    
    embeddings = generate_embeddings(text)
    embedding_path = os.path.join(EMBEDDING_DIR, f"{custom_name}_embeddings.npy")
    np.save(embedding_path, embeddings)
    
    # 生成唯一ID（使用自定义名称哈希）
    unique_id = hash(custom_name) & 0x7FFFFFFF
    index.add_with_ids(embeddings, np.array([unique_id], dtype=np.int64))
    file_registry[custom_name] = {
        "path": file_path,
        "embedding": embedding_path,
        "index_ids": [unique_id]
    }
    return gr.update(choices=list(file_registry.keys()), value=[])

# 保持其他函数定义不变
def delete_file(selected_files):
    try:
        # 添加删除进度反馈
        yield "开始删除文件..."
        for name in selected_files:
            if name in file_registry:
                # 添加调试信息
                print(f"正在删除：{name}, ID列表：{file_registry[name]['index_ids']}")
                index.remove_ids(np.array(file_registry[name]["index_ids"], dtype=np.int64))
                os.remove(file_registry[name]["path"])
                os.remove(file_registry[name]["embedding"])
                del file_registry[name]
                yield f"已删除：{name}"
        # 强制返回更新后的空值选择
        yield gr.update(choices=list(file_registry.keys()), value=[])
    except Exception as e:
        raise gr.Error(f"删除失败：{str(e)}")

def get_answer(question, selected_files):
    contexts = []
    for name in selected_files:
        full_text = read_file_content(file_registry[name]["path"])
        query_emb = generate_embeddings([question])[0]
        D, I = index.search(np.array([query_emb]), 3)
        
        paragraphs = full_text.split('\n')
        contexts.extend([f"【{name}】{paragraphs[int(idx)]}" for idx in I[0] if idx < len(paragraphs)])
    
    prompt = f"""你是一位专业法律顾问，请根据以下法律条款回答问题：
问题：{question}

相关法律条款：
{'\n'.join(contexts) if contexts else '无相关条款'}

请用中文分点给出：
1. 法律依据（引用具体条款）
2. 条款解释
3. 实务建议"""
    
    try:
        from zhipuai import ZhipuAI
        client = ZhipuAI(api_key=os.getenv('ZHIPU_API_KEY'))
        response = client.chat.completions.create(
            model="glm-4",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"生成回答时出错：{str(e)}"

# 在创建Gradio界面前初始化
init_system()

# 简化界面初始化逻辑
with gr.Blocks() as demo:
    file_list = gr.Dropdown(
        label="已上传文件",
        multiselect=True,
        choices=list(file_registry.keys()),
        interactive=True
    )

    with gr.Tab("文件管理"):
        with gr.Row():
            file_input = gr.File(label="上传法律文件", file_types=[".txt", ".docx", ".pdf"])
            custom_name = gr.Textbox(label="自定义文件标识")
            upload_btn = gr.Button("上传")
        delete_btn = gr.Button("删除选定文件")
        
        # 文件管理相关的事件绑定
        upload_btn.click(
            process_upload,
            inputs=[file_input, custom_name],
            outputs=file_list
        )
        delete_btn.click(
            delete_file,
            inputs=[file_list],
            outputs=file_list
        )

    with gr.Tab("智能问答"):
        question_input = gr.Textbox(label="输入法律问题")
        answer_output = gr.Textbox(label="法律分析建议", interactive=False)
        submit_btn = gr.Button("提交查询")
        
        # 智能问答相关的事件绑定
        submit_btn.click(
            get_answer,
            inputs=[question_input, file_list],
            outputs=answer_output
        )


# 在初始化部分添加数据库加载功能
def load_existing_files():
    for f in os.listdir(UPLOAD_DIR):
        if f.endswith(('.txt', '.docx', '.pdf')):
            name = os.path.splitext(f)[0]
            file_path = os.path.join(UPLOAD_DIR, f)
            embed_path = os.path.join(EMBEDDING_DIR, f"{name}_embeddings.npy")
            
            if os.path.exists(embed_path):
                # 加载到文件注册表和FAISS
                embeddings = np.load(embed_path)
                save_to_faiss(embeddings, index)
                file_registry[name] = {
                    "path": file_path,
                    "embedding": embed_path,
                    "index_id": len(file_registry)
                }

# 在文件末尾添加路径验证（main块中）
if __name__ == "__main__":
    print(f"存储路径验证：\n{UPLOAD_DIR}\n{EMBEDDING_DIR}")
    demo.launch(server_port=7860)
